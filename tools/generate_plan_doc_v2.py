from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "计划书"
OUT_PATH = OUT_DIR / "ITSM客服工单系统产品计划书_v2.0.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
FILL = "F2F4F7"


def set_font(run, size=10.5, color=INK, bold=False):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(key), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill=FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell_text(t.rows[0].cells[i], header, bold=True)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], str(value))
    doc.add_paragraph()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13, color=BLUE if level <= 2 else DARK, bold=True)


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.2)


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    header = sec.header.paragraphs[0]
    header.text = "ITSM 客服工单系统产品计划书 v2.0"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(32)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("ITSM 客服工单系统产品计划书")
    set_font(r, size=24, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("第二版：Java 人工处理链路与接口交付闭环")
    set_font(r, size=13, color=DARK)

    table(doc, ["项目", "内容"], [
        ["文档版本", "v2.0"],
        ["创建日期", "2026-08-24"],
        ["来源", "产品经理 Agent 基于 Java 第一版评估生成"],
        ["本版重点", "转人工、客服队列、受理、技术分析、接口交付闭环"],
    ])

    heading(doc, "1. 本版变化摘要")
    for item in [
        "将 Java 第一版的用户提问与 Agent 预处理确认为可继续推进的基础版本。",
        "将第 2 轮重点调整为人工处理链路前半段：转人工、客服队列、受理、技术分析。",
        "明确 Python Agent 继续保留为问答、摘要、建议能力，不直接修改工单核心状态。",
        "把团队老大、前端、后端、测试、分析师、file-agent 的接口交付闭环纳入计划书。",
    ]:
        bullet(doc, item)

    heading(doc, "2. 第二版产品目标")
    para(doc, "第二版目标是在 Java 主系统内完成人工客服接管能力，让 Agent 无法解决或用户主动转人工的问题，能够稳定进入客服队列，并由客服完成受理和技术分析。")

    heading(doc, "3. 本版接口任务")
    table(doc, ["轮次", "接口范围", "完成标准"], [
        ["第1轮", "用户提交问题、Agent 预处理、工单详情、Agent 能力预留", "已完成 Java 第一版并通过评分"],
        ["第2轮", "用户转人工、客服队列、客服受理、技术分析", "状态机、租户、角色校验全部落地"],
        ["第3轮", "技术支持、解决、用户确认、评价、关闭、重开", "人工处理闭环完成"],
    ])

    heading(doc, "4. 关键约束")
    for item in [
        "接口文档是前端 Agent、后端 Agent、测试 Agent 的唯一行动依据。",
        "前端 Agent 不得定义接口、字段、错误码、状态机或后端规则。",
        "后端 Agent 不得定义页面交互、按钮状态或前端展示规则。",
        "测试通过、分析师评分通过、file-agent 归档完成后，才能重置前端、后端、测试 Agent 会话。",
        "重置后必须基于下一份接口文档开展新任务，不沿用上一轮临时上下文。",
    ]:
        bullet(doc, item)

    heading(doc, "5. 后续待办")
    for item in [
        "将当前内存工单服务替换为 MySQL Repository。",
        "补充 SLA 实例、合同绑定和超时通知事件。",
        "把 Python Agent Stub 替换为真实 Python 服务调用。",
        "补充接口幂等键、Outbox 事件表和 MQ 重试策略。",
    ]:
        bullet(doc, item)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
