from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "计划书"
OUT_PATH = OUT_DIR / "ITSM客服工单系统产品计划书_v1.0.docx"

PRIMARY_BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=9.4)
        set_cell_shading(table.rows[0].cells[index], LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value), size=8.8)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    size = 16 if level == 1 else 13 if level == 2 else 12
    color = PRIMARY_BLUE if level in (1, 2) else DARK_BLUE
    for run in paragraph.runs:
        set_east_asia_font(run)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = True
    return paragraph


def add_para(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_east_asia_font(first)
        first.font.size = Pt(10.5)
        first.font.color.rgb = RGBColor.from_string(INK)
        first.bold = True
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_east_asia_font(rest)
        rest.font.size = Pt(10.5)
        rest.font.color.rgb = RGBColor.from_string(INK)
    else:
        run = paragraph.add_run(text)
        set_east_asia_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""

    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    set_east_asia_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_run = body_para.add_run(body)
    set_east_asia_font(body_run)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor.from_string(INK)

    set_table_geometry(table, [9360])
    doc.add_paragraph()


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, PRIMARY_BLUE, 16, 8),
        ("Heading 2", 13, PRIMARY_BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "ITSM 客服工单系统产品计划书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asia_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "版本 v1.0 | 产品经理 Agent 维护"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_east_asia_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("ITSM 客服工单系统产品计划书")
    set_east_asia_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string(PRIMARY_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("第一版：产品架构、核心流程与一期范围")
    set_east_asia_font(subtitle_run)
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    add_table(doc, ["项目", "内容"], [
        ["文档版本", "v1.0"],
        ["创建日期", "2026-08-24"],
        ["维护角色", "产品经理 Agent"],
        ["适用对象", "产品负责人、架构师、Java 后端研发、Python Agent 研发、测试、运维、业务评审方"],
        ["文档定位", "初步产品计划书，聚焦服务拆分、核心业务闭环、数据模型和验收边界"],
    ], [1800, 7560])

    add_callout(
        doc,
        "本版摘要",
        "本版将 ITSM 客服工单系统定义为内外兼容的企业级服务平台。一期优先完成用户提问、Agent 回答、转人工、客服受理、技术分析、处理、SLA 约束、解决确认、评价和审计留痕的闭环。",
    )

    doc.add_page_break()

    add_heading(doc, "1. 产品定位与目标", 1)
    add_para(doc, "产品定位：面向内部员工与外部客户的 ITSM 客服工单系统，优先解决电脑桌面、办公软件、账号权限、网络/VPN、终端故障等高频 IT 支持问题。", "产品定位：")
    add_para(doc, "业务目标：用智能客服 Agent 承接可标准化回答的问题，用人工客服闭环复杂问题，并通过 SLA、审计和报表保证服务质量可追踪。", "业务目标：")
    add_para(doc, "技术目标：Java 微服务群作为 ITSM 业务事实源，Python Agent 独立承载知识库问答、检索增强、摘要生成和转人工建议能力。", "技术目标：")

    add_heading(doc, "1.1 一期成功标准", 2)
    for item in [
        "用户可以在统一入口提交问题并获得 Agent 初步回答。",
        "Agent 低置信度、超时、异常或用户主动要求人工时，可以自动转入人工工单。",
        "客服端可以完成受理、技术分析、技术支持、状态变更、解决、关闭和重开。",
        "外部客户工单可以绑定合同与 SLA 策略，内部员工走默认服务策略。",
        "系统可以记录关键操作、状态变更、Agent 回答、附件和 SLA 事件。",
        "报表可以统计工单量、解决率、SLA 达成率、Agent 命中率和用户满意度。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 用户与典型场景", 1)
    add_table(doc, ["用户类型", "典型诉求", "关键能力"], [
        ["内部员工", "电脑无法开机、Office 异常、VPN 连接失败、账号登录失败、打印机不可用", "快速提问、查看答案、补充截图、转人工、确认解决、评价"],
        ["外部客户", "客户侧系统使用问题、终端环境问题、合同内服务请求、响应时限关注", "客户身份识别、合同/SLA 绑定、工单进度透明、服务评价"],
        ["一线客服", "查看待处理工单、受理、分派、沟通、记录处理过程", "队列管理、上下文查看、状态流转、内部备注、通知"],
        ["技术专家", "处理复杂技术问题、分析日志、提供解决方案", "协同处理、技术分析记录、解决方案沉淀到知识库"],
        ["管理员/主管", "管理组织、权限、SLA、知识库、报表与质量", "多租户管理、权限配置、指标报表、审计追踪"],
    ], [1500, 4300, 3560])

    add_heading(doc, "3. 一期范围与暂缓范围", 1)
    add_heading(doc, "3.1 一期范围", 2)
    for item in [
        "用户侧：提交问题、查看 Agent 回答、转人工、上传附件、补充信息、查看进度、确认解决、评价。",
        "客服端：工单池、受理、分派、技术分析、技术支持、状态修改、解决方案、关闭与重开。",
        "Agent：知识库问答、置信度评估、来源引用、转人工判断、工单摘要、建议分类和优先级。",
        "SLA/合同：合同绑定、响应时限、解决时限、临近超时提醒、超时升级事件。",
        "基础治理：权限、多租户隔离、附件、通知、审计、基础报表。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 暂缓范围", 2)
    for item in [
        "自动远程控制用户电脑。",
        "自动执行修复脚本或高权限诊断脚本。",
        "终端资产扫描与完整 CMDB。",
        "复杂 ITIL 流程，如变更管理、问题管理、发布管理。",
        "多渠道全量接入，如电话客服、公众号、小程序、第三方客服系统深度集成。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 整体架构方案", 1)
    add_para(doc, "系统采用“网关 + Java 微服务群 + Python Agent 服务 + 企业中间件”的完整微服务方向。Java 主系统负责工单与业务数据的最终一致性，Python Agent 通过受控 API 返回回答、摘要和建议，不直接修改核心业务状态。")

    add_heading(doc, "4.1 服务拆分", 2)
    add_table(doc, ["服务", "职责边界", "核心数据归属"], [
        ["API 网关", "统一入口、认证、租户识别、限流、路由、审计入口", "无核心业务数据"],
        ["认证与租户服务", "账号、角色、权限、组织、租户隔离", "tenant、organization、user_account、role、permission"],
        ["用户/客户服务", "内部员工、外部客户、客户组织、联系方式、合同归属", "customer_profile、support_agent"],
        ["工单服务", "工单创建、状态机、分类、优先级、分派、关闭、重开", "ticket、ticket_status_history、ticket_assignment"],
        ["会话服务", "用户问答会话、Agent 消息、人工消息、转人工上下文", "conversation_session、message、internal_note"],
        ["客服工作台服务", "待办队列、受理池、技术分析、协同处理、客服操作聚合", "工作台视图数据，业务状态回写工单服务"],
        ["SLA/合同服务", "合同、服务目录、响应/解决时限、预警、升级规则", "contract、service_catalog、sla_policy、sla_instance"],
        ["知识库服务", "FAQ、知识文章、版本、审核、发布、检索索引管理", "knowledge_article、knowledge_version、knowledge_chunk"],
        ["Agent 编排服务", "调用 Python Agent、超时处理、置信度判断、转人工策略", "agent_session、agent_answer、handoff_request"],
        ["Python Agent 服务", "问题理解、知识检索、LLM 回答、来源引用、摘要生成", "不持有业务事实数据，只保留必要运行日志"],
        ["附件服务", "截图、日志、录屏、文档附件的上传下载和权限校验", "message_attachment、对象存储引用"],
        ["通知服务", "站内信、邮件、企业 IM 扩展、SLA 提醒、状态通知", "notification_task"],
        ["审计与报表服务", "审计日志、操作事件、工单趋势、客服绩效、Agent 命中率", "audit_log、operation_event、报表宽表"],
    ], [1500, 4400, 3460])

    add_heading(doc, "4.2 企业中间件选型", 2)
    add_table(doc, ["能力", "默认选型", "用途"], [
        ["注册配置", "Nacos", "服务注册、配置管理、环境隔离"],
        ["服务调用", "REST + OpenFeign", "Java 服务间同步调用，一期不引入 gRPC"],
        ["消息队列", "RocketMQ", "工单事件、通知事件、Agent 异步任务、SLA 预警事件"],
        ["数据库", "MySQL", "按服务拆分 schema，核心业务避免跨库强事务"],
        ["缓存", "Redis", "会话缓存、Token、热点知识、限流、临时状态"],
        ["搜索", "OpenSearch/Elasticsearch", "工单检索、知识库全文检索、日志检索"],
        ["向量库", "Milvus", "知识库 Embedding 检索，原文仍由知识库服务管理"],
        ["对象存储", "MinIO/S3", "截图、日志、录屏、文档附件"],
        ["可观测性", "Prometheus + Grafana + OpenTelemetry", "指标、链路追踪、服务健康监控"],
    ], [1600, 2500, 5260])

    add_heading(doc, "5. 核心业务流程", 1)
    for item in [
        "用户在用户侧入口提交问题，系统创建会话并记录用户身份、租户、问题文本和附件。",
        "会话服务将问题交给 Agent 编排服务，Agent 编排服务调用 Python Agent。",
        "Python Agent 进行知识库检索与回答生成，返回回答内容、置信度、来源引用、建议分类、建议优先级和是否建议转人工。",
        "如果置信度达标，用户侧展示答案；用户点击“已解决”后，系统记录为 Agent 解决。",
        "如果置信度不足、Agent 超时/异常，或用户主动点击“转人工”，系统创建正式工单。",
        "转人工时系统自动生成工单摘要，包含原始问题、Agent 回答、已尝试方案、用户环境、建议分类和推荐优先级。",
        "客服在工作台受理工单，进入技术分析，必要时分派给技术专家协同处理。",
        "客服提交解决方案后，工单进入待用户确认；用户确认后进入已解决，再按规则关闭。",
        "SLA 服务持续计算响应时限和解决时限，临近超时和已超时均触发通知及升级事件。",
        "所有消息、附件、状态变更、Agent 回答、客服操作、SLA 事件写入审计与历史记录。",
    ]:
        add_number(doc, item)

    add_heading(doc, "6. 工单状态机", 1)
    add_table(doc, ["状态", "含义", "主要进入条件"], [
        ["SUBMITTED", "用户已提交", "用户提交问题后生成初始记录"],
        ["AGENT_PROCESSING", "Agent 正在回答", "会话进入智能客服处理"],
        ["AGENT_ANSWERED", "Agent 已回答，等待用户确认", "Agent 高置信度返回答案"],
        ["PENDING_HUMAN", "等待人工处理", "用户转人工、Agent 低置信度、Agent 超时或异常"],
        ["ACCEPTED", "客服已受理", "客服从队列中受理工单"],
        ["TECH_ANALYSIS", "技术分析中", "客服或技术专家开始诊断"],
        ["IN_SUPPORT", "技术支持处理中", "已有处理方案或正在远程指导用户"],
        ["PENDING_USER_CONFIRM", "等待用户确认", "客服提交解决方案"],
        ["RESOLVED", "已解决", "用户确认解决或客服按规则标记解决"],
        ["CLOSED", "已关闭", "解决后超过确认期或用户确认关闭"],
        ["REOPENED", "已重开", "用户反馈未解决或问题复现"],
        ["CANCELLED", "已取消", "用户取消或客服判定无效工单"],
    ], [2200, 3300, 3860])
    add_callout(doc, "状态机约束", "状态流转必须由工单服务统一校验。非法跳转需要返回明确错误，并写入审计日志。客服工作台只能通过工单服务暴露的命令接口修改状态，不能直接更新工单表。")

    add_heading(doc, "7. 核心数据模型", 1)
    add_table(doc, ["领域", "核心表/对象", "说明"], [
        ["租户与账号", "tenant、organization、user_account、customer_profile、support_agent、role、permission", "支撑内外兼容、多组织、多角色、权限隔离"],
        ["工单主数据", "ticket、ticket_status_history、ticket_assignment、ticket_participant、ticket_tag、ticket_category", "保存工单事实、状态历史、分派和分类信息"],
        ["沟通记录", "conversation_session、message、message_attachment、internal_note", "保存用户、Agent、客服、内部协同的完整上下文"],
        ["Agent 记录", "agent_session、agent_answer、agent_source_reference、handoff_request", "保存 Agent 调用、回答、引用来源和转人工摘要"],
        ["知识库", "knowledge_article、knowledge_category、knowledge_tag、knowledge_version、knowledge_chunk、embedding_index_ref", "保存知识原文、版本、切片和向量索引引用"],
        ["SLA/合同", "contract、service_catalog、sla_policy、sla_instance、sla_event、escalation_rule", "支撑外部客户合同和内外部服务时限策略"],
        ["审计通知", "audit_log、notification_task、operation_event", "保存操作审计、通知任务和异步事件"],
    ], [1700, 4100, 3560])

    add_heading(doc, "8. 接口边界", 1)
    sections = [
        ("8.1 用户侧 API", ["提交问题", "查看 Agent 回答", "转人工", "上传附件", "查看工单", "补充信息", "确认解决", "评价服务"]),
        ("8.2 客服端 API", ["查询待办队列", "受理工单", "分派工单", "修改状态", "添加技术分析记录", "提交解决方案", "关闭/重开工单"]),
        ("8.3 Agent API", ["提交问题给 Python Agent", "获取回答与置信度", "回传知识来源", "生成转人工摘要", "返回建议分类和建议优先级"]),
        ("8.4 领域事件", ["工单创建", "状态变更", "转人工", "客服受理", "SLA 即将超时", "SLA 已超时", "工单解决", "用户评价"]),
    ]
    for heading, items in sections:
        add_heading(doc, heading, 2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "9. Agent 产品边界", 1)
    add_para(doc, "一期 Agent 定位为“知识库问答 + 转人工辅助”。Agent 可以理解问题、检索知识、生成回答、输出来源、评估置信度并生成转人工摘要，但不能直接关闭工单、不能直接修改核心状态、不能远程控制用户电脑、不能执行修复脚本。")
    add_table(doc, ["场景", "Agent 行为", "系统降级"], [
        ["知识库命中且置信度高", "直接回答并展示来源引用", "用户可确认解决或转人工"],
        ["知识库命中但置信度低", "给出谨慎回答并建议转人工", "自动创建待人工工单"],
        ["无知识库命中", "生成问题摘要和建议分类", "转人工，由客服处理并沉淀知识"],
        ["Agent 服务超时/异常", "Java 编排服务记录失败原因", "直接转人工，不阻塞用户提交"],
    ], [2300, 3600, 3460])

    add_heading(doc, "10. SLA、权限与审计", 1)
    for heading, items in [
        ("10.1 SLA 策略", [
            "外部客户优先按合同和服务目录绑定 SLA 策略。",
            "内部员工按默认服务策略计算响应时限和解决时限。",
            "SLA 从问题提交或转人工节点开始计算，具体起算点需要在接口设计阶段确认。",
            "临近超时触发客服提醒，已超时触发主管升级事件。",
        ]),
        ("10.2 权限与租户隔离", [
            "所有核心接口必须校验租户、用户身份、角色权限和数据归属。",
            "外部客户只能访问本客户组织下的工单、会话和附件。",
            "客服可按技能组、队列和授权范围访问工单。",
            "管理员可以配置组织、角色、SLA、服务目录和知识库审核流程。",
        ]),
        ("10.3 审计要求", [
            "状态变更、分派、受理、解决、关闭、重开必须记录操作人、时间、来源和原因。",
            "Agent 回答、来源引用、置信度、转人工摘要必须可追溯。",
            "附件上传、下载、删除需要记录审计日志。",
            "SLA 预警、超时、升级事件需要进入审计与报表。",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "11. 验收标准", 1)
    add_table(doc, ["编号", "验收场景", "通过标准"], [
        ["A01", "Agent 成功回答", "用户提交问题后收到回答，点击已解决后系统记录 Agent 解决结果"],
        ["A02", "Agent 失败转人工", "低置信度、超时或异常时自动创建人工工单并生成摘要"],
        ["A03", "用户主动转人工", "客服端可查看原始问题、Agent 回答、附件和上下文"],
        ["A04", "客服状态流转", "客服按状态机处理，非法跳转被拒绝且记录审计"],
        ["A05", "SLA 计算", "人工工单能正确计算响应/解决时限并触发提醒"],
        ["A06", "内外兼容", "外部客户绑定合同 SLA，内部员工走默认服务策略"],
        ["A07", "租户隔离", "外部客户不能访问其他客户或内部组织数据"],
        ["A08", "知识库引用", "知识文章发布后 Agent 可检索并记录来源引用"],
        ["A09", "附件权限", "附件可上传下载，且按租户和角色校验权限"],
        ["A10", "审计留痕", "客服操作、状态变更、Agent 回答、SLA 事件均有审计"],
        ["A11", "降级处理", "MQ 重试、Agent 不可用、检索失败时可降级为人工工单"],
        ["A12", "报表统计", "可统计工单量、解决率、SLA 达成率、Agent 命中率和满意度"],
    ], [900, 3100, 5360])

    add_heading(doc, "12. 非功能需求", 1)
    for item in [
        "安全：核心接口具备鉴权、租户校验、参数校验和统一错误响应。",
        "可靠性：服务间调用具备超时、重试、熔断和告警策略。",
        "一致性：工单创建、状态变更、转人工事件具备幂等能力，跨服务协作用事件驱动和 Outbox 保证最终一致。",
        "可观测性：关键链路通过 traceId 串联用户提问、Agent 调用、工单创建和客服处理全过程。",
        "可恢复性：知识库、工单、附件、审计数据具备基础备份与恢复策略。",
        "可扩展性：预留 CMDB、远程诊断、多渠道接入、自动修复脚本和完整 ITIL 流程扩展点。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. 待确认事项", 1)
    add_table(doc, ["待确认项", "影响范围", "建议处理时间"], [
        ["SLA 起算点是用户提交问题还是确认转人工", "SLA 计算、报表、合同条款", "接口设计前"],
        ["Agent 置信度阈值和人工兜底规则", "Agent 编排、用户体验、客服工单量", "Agent 联调前"],
        ["外部客户合同与服务目录的最小字段", "客户服务、SLA、权限", "数据库设计前"],
        ["知识库审核流程是否需要多级审批", "知识库服务、客服沉淀流程", "知识库一期设计前"],
        ["报表首期是否需要主管个人绩效指标", "报表服务、权限、组织考核", "管理端设计前"],
    ], [3000, 3300, 3060])

    add_heading(doc, "14. 版本管理规则", 1)
    for item in [
        "所有计划书版本统一保存在项目根目录的“计划书”文件夹。",
        "文件命名采用“ITSM客服工单系统产品计划书_v版本号.docx”。",
        "每次优化必须基于最新版本生成新文档，不覆盖历史版本。",
        "每个版本必须在首页记录版本号、日期、维护角色和本版摘要。",
        "产品经理 Agent 负责维护版本连续性、变更摘要和待确认事项。",
    ]:
        add_bullet(doc, item)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_east_asia_font(run)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
